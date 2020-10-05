from os.path import join, exists
from docx import *
import docx
from os import listdir, mkdir
from os import walk
import subprocess
import re

path = {
    "zipFolder": "arxiv",
    "folder": "arxiv_2020_9"
}
keywords = ["3d", "point cloud", "point set", "mesh"]

if not exists(path["folder"]):
    mkdir(path["folder"])
    for zipFile in listdir(path["zipFolder"]):
        if zipFile.endswith("zip"):
            subprocess.run(["unzip",
                            join(path["zipFolder"], zipFile),
                            "-d",
                            join(path["folder"], zipFile[:-4])])

def add_hyperlink(paragraph, url, text):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


class DocWriter:

    def __init__(self):
        self.document = Document()

    def write(self, fileName, papers):
        self.document.add_heading(fileName.split('/')[-1], 0)
        for keyword in keywords:
            self.document.add_heading(keyword, level=1)
            if len(papers[keyword]) == 0:
                self.document.add_paragraph("None")
            for idx, each in enumerate(papers[keyword]):
                p = self.document.add_paragraph()
                p.add_run(str(idx+1)+'.\n').bold = True
                for key, value in each.items():

                    p.add_run(key+": ").bold = True

                    if key == "Address":
                        add_hyperlink(p, value, "arXiv:"+value.split("/")[-1])
                    else:
                        p.add_run("None\n" if value is None else value+"\n")

    def save(self):
        self.document.save('RelatedPapers.docx')


formWriter = DocWriter()


def matchPaper(paper, keyword):
    if re.search(keyword, paper) is not None:
        title = re.findall("title: .+\n", paper)[0][len("title: "):-2]
        authors = re.findall("authors: .+\n", paper)[0][len("authors: "):-2]
        categories = re.findall("categories: .+\n", paper)[0][len("categories: "):-2]
        comments = re.findall("comments: .+\n", paper)
        abstract = re.findall("\\\\.+\\\\", paper, flags=re.DOTALL)
        link = re.findall("https://arxiv.org/abs/[0-9]+.[0-9]+", paper)[0]
        return {"Title": title,
                "Authors": authors,
                "Categories": categories,
                "Comments": comments[0] if len(comments) else None,
                "Abstract": abstract[0][len("\\\n"):-len("\n\\")] if len(abstract) else None,
                "Address": link}
    else:
        return None


def match(filePath):
    with open(filePath, "r") as file:
        content = file.read()
        result = re.findall("arXiv:[0-9]+.[0-9]+.*https://arxiv.org/abs/[0-9]+.[0-9]+.*kb", content, flags=re.DOTALL)
        result = re.split("-------+\n\\\\\\\\", result[0], flags=re.DOTALL)
        relativePaper = {}
        for keyword in keywords:
            relativePaper[keyword] = []
        for paper in result:
            for keyword in keywords:
                matchResult = matchPaper(paper.lower(), keyword)
                if matchResult is not None:
                    relativePaper[keyword].append(matchResult)

        formWriter.write(filePath, relativePaper)

for (dir_path1, dir_names, file_names) in walk(path["folder"]):
    for each in file_names:
        if each.endswith("eml"):
            filePath = join(dir_path1, each)
            match(filePath)

formWriter.save()